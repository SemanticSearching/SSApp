from docx import Document

document = Document('../static/docxs/Banco Pop Exhibit - Required Insurance BP v. 12.14.18.docx')

for para in document.paragraphs:
    print(para.text)

# for content in document.paragraphs:
#     if content.style.name == 'Heading 1' or content.style.name == 'Heading 2' or content.style.name == 'Heading 3':
#         print(content.text)
#         # print(content.txt)
#         print("next paragraph\n")

# for image in document.inline_shapes:
#     print(image.width, image.height)
#
# for table in document.tables:
#     for row in table.rows:
#       for cell in row.cells:
#         for para in cell.paragraphs:
#           print(para.text)

